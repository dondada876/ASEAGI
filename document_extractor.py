#!/usr/bin/env python3
"""
Document Extractor & Repository Builder
Extracts text from RTF, DOC, DOCX, PDF, TXT and stores in structured schema
Uses open-source Python libraries only
"""

import os
import re
import hashlib
from pathlib import Path
from typing import Optional, Dict, List, Any
from dataclasses import dataclass, asdict
from datetime import datetime
import json

# Document parsing libraries
try:
    from striprtf.striprtf import rtf_to_text  # RTF extraction
except ImportError:
    print("[WARN] Installing striprtf library...")
    os.system("pip install striprtf")
    from striprtf.striprtf import rtf_to_text

try:
    import docx  # DOCX extraction
except ImportError:
    print("[WARN] Installing python-docx library...")
    os.system("pip install python-docx")
    import docx

try:
    import PyPDF2  # PDF extraction
except ImportError:
    print("[WARN] Installing PyPDF2 library...")
    os.system("pip install PyPDF2")
    import PyPDF2

try:
    import textract  # Universal document extractor (optional, fallback)
except ImportError:
    textract = None
    print("[WARN] textract not installed (optional). Install with: pip install textract")


@dataclass
class DocumentMetadata:
    """Structured metadata for extracted documents"""
    file_path: str
    file_name: str
    file_type: str  # rtf, doc, docx, pdf, txt
    file_size: int  # bytes
    file_hash: str  # MD5 hash for deduplication
    extraction_date: str
    extraction_method: str  # Which library was used
    title: Optional[str] = None
    author: Optional[str] = None
    created_date: Optional[str] = None
    modified_date: Optional[str] = None
    page_count: Optional[int] = None
    word_count: Optional[int] = None
    char_count: Optional[int] = None

@dataclass
class ExtractedDocument:
    """Complete extracted document with content and metadata"""
    metadata: DocumentMetadata
    content: str
    content_preview: str  # First 500 chars
    sections: List[Dict[str, str]]  # Structured sections
    extraction_success: bool
    extraction_error: Optional[str] = None


class DocumentExtractor:
    """Extract text from various document formats"""

    def __init__(self, output_dir: str = "document_repository"):
        """Initialize extractor with output directory"""
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)

        # Create subdirectories
        (self.output_dir / "raw_text").mkdir(exist_ok=True)
        (self.output_dir / "metadata").mkdir(exist_ok=True)
        (self.output_dir / "json").mkdir(exist_ok=True)

    def calculate_hash(self, file_path: str) -> str:
        """Calculate MD5 hash of file for deduplication"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()

    def extract_rtf(self, file_path: str) -> tuple[str, str]:
        """Extract text from RTF file"""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()

            # Use striprtf library
            text = rtf_to_text(rtf_content)
            return text, "striprtf"

        except Exception as e:
            # Fallback: manual RTF stripping (basic)
            try:
                text = self._strip_rtf_manual(rtf_content)
                return text, "manual_rtf_strip"
            except:
                raise Exception(f"RTF extraction failed: {e}")

    def _strip_rtf_manual(self, rtf_content: str) -> str:
        """Manual RTF stripping (fallback method)"""
        # Remove RTF header and font table
        text = re.sub(r'\\rtf1.*?\\fonttbl.*?\}', '', rtf_content, flags=re.DOTALL)
        text = re.sub(r'\\colortbl;.*?\}', '', text, flags=re.DOTALL)

        # Remove formatting commands
        text = re.sub(r'\{\\[a-z]+.*?\}', '', text)
        text = re.sub(r'\\[a-z]+\d*\s*', '', text)

        # Remove braces
        text = text.replace('{', '').replace('}', '')

        # Clean up whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)

        return text.strip()

    def extract_docx(self, file_path: str) -> tuple[str, str]:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)

        # Extract all paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)

        text = '\n\n'.join(paragraphs)
        return text, "python-docx"

    def extract_pdf(self, file_path: str) -> tuple[str, str]:
        """Extract text from PDF file"""
        text_parts = []

        with open(file_path, 'rb') as f:
            pdf_reader = PyPDF2.PdfReader(f)

            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(f"--- Page {page_num + 1} ---\n{page_text}")
                except Exception as e:
                    text_parts.append(f"--- Page {page_num + 1} ---\n[Extraction Error: {e}]")

        text = '\n\n'.join(text_parts)
        return text, "PyPDF2"

    def extract_doc(self, file_path: str) -> tuple[str, str]:
        """Extract text from legacy DOC file (requires textract or antiword)"""

        # Try textract first (universal extractor)
        if textract:
            try:
                text = textract.process(file_path).decode('utf-8')
                return text, "textract"
            except Exception as e:
                print(f"⚠️  textract failed: {e}")

        # Try antiword (Linux/Mac)
        try:
            import subprocess
            result = subprocess.run(['antiword', file_path],
                                    capture_output=True, text=True, timeout=30)
            if result.returncode == 0:
                return result.stdout, "antiword"
        except:
            pass

        # Fallback: Try python-docx (might work for some .doc files)
        try:
            text, method = self.extract_docx(file_path)
            return text, f"{method}_legacy_mode"
        except:
            pass

        raise Exception("DOC extraction failed. Install 'textract' or 'antiword' for .doc support")

    def extract_txt(self, file_path: str) -> tuple[str, str]:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return text, "direct_read"

    def extract_document(self, file_path: str) -> ExtractedDocument:
        """Extract document with full metadata"""

        file_path_obj = Path(file_path)

        # Get file metadata
        file_stat = file_path_obj.stat()
        file_ext = file_path_obj.suffix.lower()

        metadata = DocumentMetadata(
            file_path=str(file_path_obj.absolute()),
            file_name=file_path_obj.name,
            file_type=file_ext[1:] if file_ext else 'unknown',
            file_size=file_stat.st_size,
            file_hash=self.calculate_hash(file_path),
            extraction_date=datetime.now().isoformat(),
            extraction_method="",
            created_date=datetime.fromtimestamp(file_stat.st_ctime).isoformat(),
            modified_date=datetime.fromtimestamp(file_stat.st_mtime).isoformat()
        )

        try:
            # Extract based on file type
            if file_ext == '.rtf':
                content, method = self.extract_rtf(file_path)
            elif file_ext == '.docx':
                content, method = self.extract_docx(file_path)
            elif file_ext == '.doc':
                content, method = self.extract_doc(file_path)
            elif file_ext == '.pdf':
                content, method = self.extract_pdf(file_path)
            elif file_ext in ['.txt', '.md']:
                content, method = self.extract_txt(file_path)
            else:
                raise Exception(f"Unsupported file type: {file_ext}")

            metadata.extraction_method = method

            # Calculate text statistics
            metadata.word_count = len(content.split())
            metadata.char_count = len(content)

            # Extract title (first non-empty line)
            lines = [l.strip() for l in content.split('\n') if l.strip()]
            if lines:
                metadata.title = lines[0][:200]  # First line, max 200 chars

            # Extract sections (simple heuristic: lines starting with # or all caps)
            sections = self._extract_sections(content)

            # Create preview
            content_preview = content[:500] + "..." if len(content) > 500 else content

            return ExtractedDocument(
                metadata=metadata,
                content=content,
                content_preview=content_preview,
                sections=sections,
                extraction_success=True,
                extraction_error=None
            )

        except Exception as e:
            return ExtractedDocument(
                metadata=metadata,
                content="",
                content_preview="",
                sections=[],
                extraction_success=False,
                extraction_error=str(e)
            )

    def _extract_sections(self, content: str) -> List[Dict[str, str]]:
        """Extract sections from document (simple heuristic)"""
        sections = []
        current_section = None
        current_content = []

        for line in content.split('\n'):
            line_stripped = line.strip()

            # Detect section headers (# Markdown or ALL CAPS HEADERS)
            is_header = False
            if line_stripped.startswith('#'):
                is_header = True
                header_text = line_stripped.lstrip('#').strip()
            elif (line_stripped.isupper() and
                  len(line_stripped) > 5 and
                  len(line_stripped) < 100):
                is_header = True
                header_text = line_stripped

            if is_header:
                # Save previous section
                if current_section:
                    sections.append({
                        'title': current_section,
                        'content': '\n'.join(current_content).strip()
                    })

                # Start new section
                current_section = header_text
                current_content = []
            else:
                if line_stripped:
                    current_content.append(line)

        # Save last section
        if current_section:
            sections.append({
                'title': current_section,
                'content': '\n'.join(current_content).strip()
            })

        return sections

    def save_extracted_document(self, doc: ExtractedDocument) -> Dict[str, str]:
        """Save extracted document to repository"""

        # Generate safe filename from original
        safe_name = re.sub(r'[^\w\-_.]', '_', doc.metadata.file_name)
        base_name = Path(safe_name).stem

        # Save raw text (clean surrogates)
        text_file = self.output_dir / "raw_text" / f"{base_name}.txt"
        # Remove surrogate characters that can't be encoded in UTF-8
        clean_content = ''.join(char for char in doc.content
                               if not (0xD800 <= ord(char) <= 0xDFFF))
        with open(text_file, 'w', encoding='utf-8', errors='ignore') as f:
            f.write(clean_content)

        # Save metadata
        meta_file = self.output_dir / "metadata" / f"{base_name}_meta.json"
        with open(meta_file, 'w', encoding='utf-8') as f:
            json.dump(asdict(doc.metadata), f, indent=2)

        # Save complete document as JSON
        json_file = self.output_dir / "json" / f"{base_name}.json"
        doc_dict = {
            'metadata': asdict(doc.metadata),
            'content': clean_content,  # Use cleaned content
            'content_preview': doc.content_preview,
            'sections': doc.sections,
            'extraction_success': doc.extraction_success,
            'extraction_error': doc.extraction_error
        }
        with open(json_file, 'w', encoding='utf-8', errors='ignore') as f:
            json.dump(doc_dict, f, indent=2, ensure_ascii=False)

        return {
            'text_file': str(text_file),
            'meta_file': str(meta_file),
            'json_file': str(json_file)
        }

    def process_directory(self, directory: str, recursive: bool = True) -> List[Dict]:
        """Process all documents in directory"""

        directory_path = Path(directory)
        results = []

        # Supported extensions
        extensions = ['.rtf', '.doc', '.docx', '.pdf', '.txt', '.md']

        # Find all files
        if recursive:
            files = []
            for ext in extensions:
                files.extend(directory_path.rglob(f'*{ext}'))
        else:
            files = []
            for ext in extensions:
                files.extend(directory_path.glob(f'*{ext}'))

        print(f"Found {len(files)} documents to process")
        print()

        for i, file_path in enumerate(files, 1):
            print(f"[{i}/{len(files)}] Processing: {file_path.name}")

            # Extract document
            doc = self.extract_document(str(file_path))

            if doc.extraction_success:
                # Save to repository
                saved_files = self.save_extracted_document(doc)

                result = {
                    'file': str(file_path),
                    'status': 'success',
                    'word_count': doc.metadata.word_count,
                    'char_count': doc.metadata.char_count,
                    'method': doc.metadata.extraction_method,
                    'saved_files': saved_files
                }

                print(f"   [OK] Success - {doc.metadata.word_count} words, {doc.metadata.char_count} chars")
                print(f"   Method: {doc.metadata.extraction_method}")

            else:
                result = {
                    'file': str(file_path),
                    'status': 'failed',
                    'error': doc.extraction_error
                }
                print(f"   [FAIL] Failed: {doc.extraction_error}")

            results.append(result)
            print()

        return results

    def create_index(self, results: List[Dict]) -> str:
        """Create searchable index of all documents"""

        index_file = self.output_dir / "document_index.json"

        index = {
            'created': datetime.now().isoformat(),
            'total_documents': len(results),
            'successful': len([r for r in results if r['status'] == 'success']),
            'failed': len([r for r in results if r['status'] == 'failed']),
            'documents': results
        }

        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(index, f, indent=2)

        print(f"[INDEX] Index created: {index_file}")
        print(f"   Total: {index['total_documents']}")
        print(f"   Success: {index['successful']}")
        print(f"   Failed: {index['failed']}")

        return str(index_file)


def main():
    """Main entry point"""

    print("=" * 80)
    print("DOCUMENT EXTRACTOR & REPOSITORY BUILDER")
    print("=" * 80)
    print()
    print("Supported formats: RTF, DOC, DOCX, PDF, TXT, MD")
    print("Libraries: striprtf, python-docx, PyPDF2, textract (optional)")
    print()

    # Example: Process ASEAGI RTF files
    extractor = DocumentExtractor(output_dir="PROJ344_document_repository")

    # Get current directory
    current_dir = Path.cwd()

    print(f"Processing directory: {current_dir}")
    print()

    # Process all documents
    results = extractor.process_directory(str(current_dir), recursive=False)

    # Create index
    index_file = extractor.create_index(results)

    print()
    print("=" * 80)
    print("EXTRACTION COMPLETE")
    print("=" * 80)
    print()
    print(f"Repository location: {extractor.output_dir}")
    print(f"Index file: {index_file}")
    print()
    print("Repository structure:")
    print(f"  [DIR] {extractor.output_dir}/")
    print(f"     [DIR] raw_text/        - Plain text files")
    print(f"     [DIR] metadata/        - Document metadata (JSON)")
    print(f"     [DIR] json/            - Complete documents (JSON)")
    print(f"     [FILE] document_index.json - Searchable index")


if __name__ == "__main__":
    main()
